import os
from docx import Document
from docx.shared import Inches
import pyscreenshot as ImageGrab


docx_file_name = '/root/report/[number].docx'

document = Document()

document.add_heading('Nordea verification test', 0)


def single_ip():
    ip = '10.10.109.14'
    port = '2001,2002,2004,2005,2006,2007,2008,2009,2010,2013,2020,2021,2022,2030,2033,2034,3011,3017,3031,7007,7019,7025,10012,10024'

    document.add_paragraph(ip + ':' + port)

    bash_command = 'nmap -p [port] -Pn [ip]'

    bash_command = bash_command.replace('[ip]', ip)
    bash_command = bash_command.replace('[port]', port)
    os.system('clear')
    print bash_command
    os.system(bash_command)
    im = ImageGrab.grab()
    im_name = '/root/report/'+ ip + '_' + port + '.png'
    im.save(im_name)
    document.add_picture(im_name, width=Inches(7))


def multiple_ip():
    ip_and_ports = [('10.10.110.107', '22'), ('10.10.239.140', '22')]

    for ip_and_port in ip_and_ports:
        bash_command = 'nmap -Pn -p [port] [ip]'
        #bash_command = 'nmap -sU -pU:[port] -Pn -n --script=ntp-info [ip]'
        #bash_command = 'ssh -vv [ip]'
        ip, port = ip_and_port
        document.add_paragraph(ip + ':' + port)
        bash_command = bash_command.replace('[ip]', ip)
        bash_command = bash_command.replace('[port]', port)
        os.system('clear')
        print bash_command
        os.system(bash_command)
        im = ImageGrab.grab()
        im_name = '/root/report/'+ ip + '_' + port + '.png'
        im.save(im_name)
        document.add_picture(im_name, width=Inches(7))


def main():
    number = '12'

    multiple_ip()
    #single_ip()

    document.save(docx_file_name.replace('[number]', number))

if __name__ == '__main__':
    main()